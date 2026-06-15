"""RAG pipeline: ChromaDB vector store, document ingestion, data summaries.

Two kinds of knowledge are embedded:
1. User-uploaded context docs (PDF / DOCX / TXT / MD) - chunked and embedded.
2. Auto-generated narratives about the uploaded dataset (per-period aggregates,
   trends, % changes) so "why" questions can also retrieve facts from the data.

Embeddings use ChromaDB's default local model (all-MiniLM-L6-v2) - no API key.
"""
import uuid

import chromadb
import numpy as np
import pandas as pd

import config
from services import settings

_client = chromadb.PersistentClient(path=config.CHROMA_DB_PATH)
_collection = _client.get_or_create_collection(
    name="knowledge", metadata={"hnsw:space": "cosine"}
)

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150


# ---------- document ingestion ----------

def _extract_text(path: str, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    if name.endswith(".docx"):
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    if name.endswith((".xlsx", ".xls", ".csv")):
        if name.endswith(".csv"):
            sheets = {"data": pd.read_csv(path)}
        else:
            sheets = pd.read_excel(path, sheet_name=None)
        parts = []
        for sheet, sdf in sheets.items():
            sdf = sdf.head(1000)  # cap rows per sheet
            parts.append(f"Sheet '{sheet}' with columns: {', '.join(map(str, sdf.columns))}.")
            for _, row in sdf.iterrows():
                parts.append("; ".join(f"{c}: {v}" for c, v in row.items() if pd.notna(v)))
        return "\n".join(parts)
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _chunk(text: str) -> list[str]:
    cfg = settings.get_all()
    size, overlap = cfg["chunk_size"], cfg["chunk_overlap"]
    text = " ".join(text.split())
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        # try to break at a sentence boundary
        if end < len(text):
            dot = text.rfind(". ", start + size // 2, end)
            if dot != -1:
                end = dot + 1
        chunks.append(text[start:end].strip())
        start = max(end - overlap, start + 1)
    return [c for c in chunks if len(c) > 40]


def ingest_document(path: str, filename: str, domain: str = "general") -> int:
    text = _extract_text(path, filename)
    chunks = _chunk(text)
    if not chunks:
        return 0
    _collection.add(
        ids=[f"doc-{uuid.uuid4().hex[:10]}-{i}" for i in range(len(chunks))],
        documents=chunks,
        metadatas=[{"source": filename, "type": "document", "domain": domain} for _ in chunks],
    )
    return len(chunks)


def seed_glossaries(glossaries: dict[str, list[str]]) -> None:
    """Idempotently embed built-in domain KPI definitions (deterministic ids)."""
    for domain, entries in glossaries.items():
        _collection.upsert(
            ids=[f"gloss-{domain}-{i}" for i in range(len(entries))],
            documents=entries,
            metadatas=[
                {"source": f"{domain} glossary", "type": "glossary", "domain": domain}
                for _ in entries
            ],
        )


# ---------- auto-generated data summaries ----------

def ingest_data_summaries(df: pd.DataFrame, filename: str, domain: str = "general") -> int:
    """Embed natural-language summaries of the dataset (overall + per period)."""
    texts = [_overall_summary(df, filename)]
    texts += _period_summaries(df, filename)
    _collection.add(
        ids=[f"data-{uuid.uuid4().hex[:10]}-{i}" for i in range(len(texts))],
        documents=texts,
        metadatas=[{"source": filename, "type": "data_summary", "domain": domain} for _ in texts],
    )
    return len(texts)


def _overall_summary(df: pd.DataFrame, filename: str) -> str:
    numeric = df.select_dtypes(include=np.number)
    lines = [f"Dataset '{filename}': {df.shape[0]} rows, columns: {', '.join(df.columns)}."]
    for col in numeric.columns:
        s = numeric[col].dropna()
        if s.empty:
            continue
        lines.append(
            f"Column '{col}': total {s.sum():,.2f}, mean {s.mean():,.2f}, "
            f"min {s.min():,.2f}, max {s.max():,.2f}."
        )
    return " ".join(lines)


def _period_summaries(df: pd.DataFrame, filename: str) -> list[str]:
    date_cols = [c for c in df.columns if pd.api.types.is_datetime64_any_dtype(df[c])]
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    if not date_cols or not numeric_cols:
        return []

    date_col = date_cols[0]
    out = []
    for freq, label in [("QE", "quarter"), ("ME", "month")]:
        try:
            grouped = df.set_index(date_col)[numeric_cols].resample(freq).sum()
        except Exception:
            continue
        grouped = grouped[grouped.abs().sum(axis=1) > 0]
        if len(grouped) < 2 or len(grouped) > 60:
            continue
        for col in numeric_cols:
            prev = None
            for period, value in grouped[col].items():
                period_name = (
                    f"Q{period.quarter} {period.year}" if label == "quarter"
                    else period.strftime("%B %Y")
                )
                sentence = f"In dataset '{filename}', {col} for {period_name} was {value:,.2f}."
                if prev not in (None, 0):
                    pct = (value - prev) / abs(prev) * 100
                    direction = "increased" if pct >= 0 else "decreased (dipped)"
                    sentence += f" Compared to the previous {label}, {col} {direction} by {abs(pct):.1f}%."
                out.append(sentence)
                prev = value
    return out


# ---------- retrieval ----------

def search(query: str, top_k: int = 6, domain: str | None = None) -> list[dict]:
    """Semantic search, scoped to the active domain (plus 'general' uploads)."""
    if _collection.count() == 0:
        return []
    where = {"domain": {"$in": [domain, "general"]}} if domain else None
    res = _collection.query(
        query_texts=[query],
        n_results=min(top_k, _collection.count()),
        where=where,
    )
    return [
        {"text": doc, "source": meta.get("source"), "type": meta.get("type")}
        for doc, meta in zip(res["documents"][0], res["metadatas"][0])
    ]


def knowledge_stats() -> dict:
    return {"chunks": _collection.count()}
